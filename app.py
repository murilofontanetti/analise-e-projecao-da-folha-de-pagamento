import streamlit as st
import pdfplumber
import pandas as pd
import re
import os
import tempfile
import requests
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# CONFIGURAÇÃO E MEMÓRIA
# ============================================================
st.set_page_config(page_title="Sistema de Impacto Orçamentário - Pessoal", layout="wide")

# ============================================================
# 🔒 SISTEMA DE LOGIN E SEGURANÇA
# ============================================================
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

if not st.session_state.autenticado:
    col_vazia1, col_login, col_vazia2 = st.columns([1, 1, 1])
    
    with col_login:
        st.title("🔒 Acesso Restrito")
        st.markdown("Por favor, insira a credencial para acessar o **Sistema de Impacto Orçamentário - Pessoal**.")
        senha_digitada = st.text_input("Senha de Acesso:", type="password")
        
        if st.button("Entrar", type="primary"):
            if senha_digitada == "9394-pub":
                st.session_state.autenticado = True
                st.rerun() 
            else:
                st.error("❌ Senha incorreta. Acesso negado.")
    st.stop() 

# ============================================================
# SISTEMA PRINCIPAL E UTILITÁRIOS
# ============================================================
st.title("📊 Análise de Impacto Orçamentário - Projeção da Folha")
st.subheader("Prefeitura Municipal de Macatuba/SP")

if 'df_processado' not in st.session_state:
    st.session_state.df_processado = None

aba1, aba2 = st.tabs(["📋 Análise de Folha (Projeção)", "👤 Simulador de Servidor"])

MESES = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4,
    "Maio": 5, "Junho": 6, "Julho": 7, "Agosto": 8,
    "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
num_para_mes = {v: k for k, v in MESES.items()}

def fmt_br(v: float) -> str:
    numero_texto = f"{v:,.2f}"
    return "R$ " + numero_texto.replace(',', 'X').replace('.', ',').replace('X', '.')

def f_moeda_limpa(v: float) -> str:
    return f"{v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

# ============================================================
# 1. FUNÇÕES DE GERAÇÃO DE ARQUIVOS (ABA 1)
# ============================================================
def gerar_excel_aba1(df_exp):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_exp.to_excel(writer, index=False, sheet_name='Projeção da Folha')
    return output.getvalue()

def gerar_word_aba1(df_exp, mes):
    doc = Document()
    titulo = doc.add_heading('Prefeitura Municipal de Macatuba/SP', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Análise de Impacto Orçamentário - Projeção da Folha\nMês de Referência da Liquidação: {mes}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=len(df_exp.columns))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_exp.columns):
        hdr_cells[i].text = col
        
    col_is_currency = ['Saldo Atual', 'Liquidado', 'Proj. Sal.', 'Prov. 13º', 'Prov. Férias', 'Total Projet.', 'Saldo Final']
    
    for _, row in df_exp.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df_exp.columns):
            val = row[col_name]
            if col_name in col_is_currency:
                row_cells[i].text = f_moeda_limpa(float(val))
            else:
                row_cells[i].text = str(val)

    doc.add_paragraph('\n\n_____________________________________________________\nMurilo Fontanetti\nContador CRC 1SP322844/0-7').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def gerar_pdf(df, mes):
    def limpar_texto(texto):
        return str(texto).encode('latin-1', 'ignore').decode('latin-1')

    url_logo = "https://upload.wikimedia.org/wikipedia/commons/c/c0/Bras%C3%A3o_Macatuba.jpg"
    logo_path = "logo_macatuba_temp.jpg"
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        response = requests.get(url_logo, headers=headers, timeout=5)
        if response.status_code == 200:
            with open(logo_path, "wb") as f:
                f.write(response.content)
    except:
        pass 

    class PDFReport(FPDF):
        def header(self):
            if os.path.exists(logo_path):
                try:
                    self.image(logo_path, x=12, y=10, w=22)
                except Exception:
                    pass
            
            self.set_font("Arial", 'B', 16)
            self.set_text_color(0, 51, 102)
            self.set_xy(38, 12)
            self.cell(0, 8, limpar_texto("Prefeitura Municipal de Macatuba/SP"), ln=True, align='L')
            
            self.set_font("Arial", 'B', 13)
            self.set_text_color(50, 50, 50)
            self.set_x(38)
            self.cell(0, 6, limpar_texto("Análise de Impacto Orçamentário - Projeção da Folha"), ln=True, align='L')
            
            self.set_font("Arial", 'I', 10)
            self.set_x(38)
            self.cell(0, 6, limpar_texto(f"Mês de Referência da Liquidação: {mes}"), ln=True, align='L')
            self.ln(10)
            
            colunas = ['Cód', 'Atividade', 'Natureza', 'Saldo Atual', 'Liquidado', 'Proj. Sal.', 'Prov. 13º', 'Prov. Férias', 'Total Projet.', 'Saldo Final', 'Situação']
            larguras = [10, 46, 16, 22, 22, 20, 18, 20, 25, 25, 53] 
            
            self.set_fill_color(0, 51, 102)
            self.set_text_color(255, 255, 255)
            self.set_font("Arial", 'B', 7)
            
            for i, col in enumerate(colunas):
                self.cell(larguras[i], 7, limpar_texto(col), border=1, align='C', fill=True)
            self.ln()

    pdf = PDFReport(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    
    larguras = [10, 46, 16, 22, 22, 20, 18, 20, 25, 25, 53] 
    df_sorted = df.sort_values(by=['Código', 'Natureza Base'])
    
    current_ativ = None
    preenchimento = False 
    
    subtotals = {
        'Saldo com Reserva': 0.0, 'Liquidado': 0.0, 'Projeção Salarial': 0.0,
        'Provisão 13º': 0.0, 'Provisão Férias': 0.0, 'Total Projetado': 0.0, 'Saldo Final': 0.0
    }
    
    def reset_subtotals():
        for key in subtotals:
            subtotals[key] = 0.0

    def print_subtotals():
        pdf.set_font("Arial", 'B', 6.5)
        pdf.set_fill_color(220, 220, 220)
        pdf.set_text_color(0, 0, 0)
        
        largura_texto = larguras[0] + larguras[1] + larguras[2]
        pdf.cell(largura_texto, 6, limpar_texto("TOTAL DA ATIVIDADE:"), border=1, align='R', fill=True)
        
        pdf.cell(larguras[3], 6, limpar_texto(f_moeda_limpa(subtotals['Saldo com Reserva'])), border=1, align='R', fill=True)
        pdf.cell(larguras[4], 6, limpar_texto(f_moeda_limpa(subtotals['Liquidado'])), border=1, align='R', fill=True)
        pdf.cell(larguras[5], 6, limpar_texto(f_moeda_limpa(subtotals['Projeção Salarial'])), border=1, align='R', fill=True)
        pdf.cell(larguras[6], 6, limpar_texto(f_moeda_limpa(subtotals['Provisão 13º'])), border=1, align='R', fill=True)
        pdf.cell(larguras[7], 6, limpar_texto(f_moeda_limpa(subtotals['Provisão Férias'])), border=1, align='R', fill=True)
        pdf.cell(larguras[8], 6, limpar_texto(f_moeda_limpa(subtotals['Total Projetado'])), border=1, align='R', fill=True)
        
        if subtotals['Saldo Final'] < 0:
            pdf.set_text_color(200, 0, 0)
        else:
            pdf.set_text_color(0, 100, 0)
            
        pdf.cell(larguras[9], 6, limpar_texto(f_moeda_limpa(subtotals['Saldo Final'])), border=1, align='R', fill=True)
        pdf.set_text_color(0, 0, 0)
        
        sit_total = "Adicionar Crédito Geral" if subtotals['Saldo Final'] < 0 else "Suficiente"
        pdf.cell(larguras[10], 6, limpar_texto(sit_total), border=1, align='C', fill=True)
        pdf.ln()

    for index, row in df_sorted.iterrows():
        if current_ativ is not None and current_ativ != row['Código']:
            print_subtotals()
            preenchimento = False 
            
        if current_ativ != row['Código']:
            current_ativ = row['Código']
            reset_subtotals()

        subtotals['Saldo com Reserva'] += row['Saldo com Reserva']
        subtotals['Liquidado'] += row['Liquidado']
        subtotals['Projeção Salarial'] += row['Projeção Salarial']
        subtotals['Provisão 13º'] += row['Provisão 13º']
        subtotals['Provisão Férias'] += row['Provisão Férias']
        subtotals['Total Projetado'] += row['Total Projetado']
        subtotals['Saldo Final'] += row['Saldo Final']

        if preenchimento:
            pdf.set_fill_color(245, 245, 245) 
        else:
            pdf.set_fill_color(255, 255, 255) 
            
        pdf.set_font("Arial", '', 6.5)
        ativ_texto = limpar_texto(str(row['Atividade']))
        while pdf.get_string_width(ativ_texto) > 44:
            ativ_texto = ativ_texto[:-1]
        
        pdf.cell(larguras[0], 6, limpar_texto(row['Código']), border=1, align='C', fill=preenchimento)
        pdf.cell(larguras[1], 6, ativ_texto, border=1, fill=preenchimento)
        
        pdf.cell(larguras[2], 6, limpar_texto(row['Natureza Base']), border=1, align='C', fill=preenchimento)
        pdf.cell(larguras[3], 6, limpar_texto(f_moeda_limpa(row['Saldo com Reserva'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[4], 6, limpar_texto(f_moeda_limpa(row['Liquidado'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[5], 6, limpar_texto(f_moeda_limpa(row['Projeção Salarial'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[6], 6, limpar_texto(f_moeda_limpa(row['Provisão 13º'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[7], 6, limpar_texto(f_moeda_limpa(row['Provisão Férias'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[8], 6, limpar_texto(f_moeda_limpa(row['Total Projetado'])), border=1, align='R', fill=preenchimento)
        
        if row['Saldo Final'] < 0:
            pdf.set_text_color(200, 0, 0)
        else:
            pdf.set_text_color(0, 0, 0)
            
        pdf.cell(larguras[9], 6, limpar_texto(f_moeda_limpa(row['Saldo Final'])), border=1, align='R', fill=preenchimento)
        pdf.set_text_color(0, 0, 0) 
        
        sit = str(row['Mês do Crédito']).replace('✅', '').replace('⚠️', '').strip()
        pdf.cell(larguras[10], 6, limpar_texto(sit), border=1, align='C', fill=preenchimento)
        pdf.ln()
        preenchimento = not preenchimento 

    if current_ativ is not None:
        print_subtotals()

    if pdf.get_y() > 165:
        pdf.add_page() 

    pdf.ln(25)
    pdf.set_font("Arial", '', 9)
    pdf.cell(0, 5, limpar_texto("_____________________________________________________"), ln=True, align='C')
    pdf.set_font("Arial", 'B', 9)
    pdf.cell(0, 5, limpar_texto("Murilo Fontanetti"), ln=True, align='C')
    pdf.set_font("Arial", '', 9)
    pdf.cell(0, 5, limpar_texto("Contador CRC 1SP322844/0-7"), ln=True, align='C')
        
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf.output(tmp.name)
        with open(tmp.name, "rb") as f:
            pdf_bytes = f.read()
            
    try:
        if os.path.exists(logo_path):
            os.remove(logo_path)
    except Exception:
        pass
        
    return pdf_bytes

# ============================================================
# 2. FUNÇÕES DE GERAÇÃO DE ARQUIVOS (ABA 2 - SIMULADOR)
# ============================================================
def gerar_excel_simulador(df_comp, df_sim):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_comp.to_excel(writer, index=False, sheet_name='Quadro Comparativo')
        df_sim.to_excel(writer, index=False, sheet_name='Memória de Cálculo')
    return output.getvalue()

def gerar_word_simulador(atividade, operacao, mes_inicio, impacto_final, df_comp, df_sim, desc):
    doc = Document()
    doc.add_heading('Prefeitura Municipal de Macatuba/SP', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Simulação de Impacto Orçamentário').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Parâmetros da Simulação', level=1)
    p = doc.add_paragraph()
    p.add_run(f'Atividade Avaliada: {atividade}\n')
    operacao_limpa = operacao.replace("➕", "").replace("➖", "").replace("🔀", "").strip()
    p.add_run(f'Operação: {operacao_limpa}\n')
    p.add_run(f'Mês de Início: {mes_inicio}\n')
    if desc:
        p.add_run(f'Descrição: {desc}\n')
    p.add_run(f'Impacto Global Final: {fmt_br(impacto_final)}').bold = True
    
    doc.add_heading('2. Memória de Cálculo', level=1)
    t_sim = doc.add_table(rows=1, cols=2)
    t_sim.style = 'Table Grid'
    t_sim.rows[0].cells[0].text = 'Item'
    t_sim.rows[0].cells[1].text = 'Valor'
    for _, row in df_sim.iterrows():
        cells = t_sim.add_row().cells
        cells[0].text = str(row['Item'])
        cells[1].text = f_moeda_limpa(float(row['Valor']))
        
    doc.add_heading('3. Quadro Comparativo por Natureza', level=1)
    t_comp = doc.add_table(rows=1, cols=len(df_comp.columns))
    t_comp.style = 'Table Grid'
    for i, col in enumerate(df_comp.columns):
        t_comp.rows[0].cells[i].text = col
        
    col_is_currency_comp = ["Dotação Atual", "Projetado (Antes)", "Impacto Simulado", "Projetado (Depois)", "Saldo Final (Antes)", "Saldo Final (Depois)"]
    for _, row in df_comp.iterrows():
        cells = t_comp.add_row().cells
        for i, col_name in enumerate(df_comp.columns):
            val = row[col_name]
            if col_name in col_is_currency_comp:
                cells[i].text = f_moeda_limpa(float(val))
            else:
                cells[i].text = str(val)

    doc.add_paragraph('\n\n_____________________________________________________\nMurilo Fontanetti\nContador CRC 1SP322844/0-7').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def gerar_pdf_simulador(atividade_nome, operacao, mes_inicio, impacto_final, df_comparativo, dados_impacto, descricao_estimativa):
    def limpar_texto(texto):
        return str(texto).encode('latin-1', 'ignore').decode('latin-1')

    pdf = FPDF(orientation='L', unit='mm', format='A4') 
    pdf.add_page()
    
    url_logo = "https://upload.wikimedia.org/wikipedia/commons/c/c0/Bras%C3%A3o_Macatuba.jpg"
    logo_path = "logo_macatuba_sim_temp.jpg"
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        response = requests.get(url_logo, headers=headers, timeout=5)
        if response.status_code == 200:
            with open(logo_path, "wb") as f:
                f.write(response.content)
    except:
        pass 
        
    if os.path.exists(logo_path):
        try:
            pdf.image(logo_path, x=12, y=10, w=22)
        except Exception:
            pass
            
    pdf.set_font("Arial", 'B', 16)
    pdf.set_text_color(0, 51, 102)
    pdf.set_xy(38, 12)
    pdf.cell(0, 8, limpar_texto("Prefeitura Municipal de Macatuba/SP"), ln=True, align='L')
    
    pdf.set_font("Arial", 'B', 13)
    pdf.set_text_color(50, 50, 50)
    pdf.set_x(38)
    pdf.cell(0, 6, limpar_texto("Análise de Impacto Orçamentário - Projeção da Folha"), ln=True, align='L')
    
    pdf.set_font("Arial", 'I', 10)
    pdf.set_x(38)
    pdf.cell(0, 6, limpar_texto(f"Atividade Avaliada: {atividade_nome}"), ln=True, align='L')
    pdf.ln(12) 

    pdf.set_font("Arial", 'B', 11)
    pdf.set_fill_color(0, 51, 102)
    pdf.set_text_color(255, 255, 255) 
    pdf.cell(0, 8, limpar_texto(" 1. PARÂMETROS DA SIMULAÇÃO"), border=1, ln=True, align='L', fill=True)
    
    pdf.set_font("Arial", '', 10)
    pdf.set_text_color(0, 0, 0) 
    
    operacao_limpa = operacao.replace("➕", "").replace("➖", "").replace("🔀", "").strip()
    pdf.cell(0, 6, limpar_texto(f"  - Operação: {operacao_limpa}"), border='L,R', ln=True)
    pdf.cell(0, 6, limpar_texto(f"  - Mês de Início: {mes_inicio} (Considerando o impacto até Dezembro)"), border='L,R', ln=True)
    if descricao_estimativa:
        pdf.cell(0, 6, limpar_texto(f"  - Descrição: {descricao_estimativa}"), border='L,R', ln=True)
    
    cor_imp = (200,0,0) if impacto_final > 0 else (0,150,0)
    pdf.cell(48, 6, limpar_texto("  - Impacto Global Final: "), border='L,B')
    pdf.set_font("Arial", 'B', 10)
    pdf.set_text_color(*cor_imp)
    pdf.cell(0, 6, limpar_texto(f"R$ {f_moeda_limpa(impacto_final)}"), border='R,B', ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 9)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 6, limpar_texto("  MEMÓRIA DE CÁLCULO (Itens processados na simulação):"), border=1, ln=True, fill=True)
    pdf.set_font("Arial", '', 9)
    for index, row in dados_impacto.iterrows():
        pdf.cell(140, 5, limpar_texto(f"    - {row['Item']}"), border='L,R')
        pdf.cell(0, 5, limpar_texto(f"R$ {f_moeda_limpa(row['Valor'])}"), border='R', ln=True, align='R')
    pdf.cell(0, 2, limpar_texto(""), border='L,R,B', ln=True) 
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 11)
    pdf.set_fill_color(0, 51, 102)
    pdf.set_text_color(255, 255, 255) 
    pdf.cell(0, 8, limpar_texto(" 2. QUADRO COMPARATIVO POR NATUREZA DE DESPESA (ANTES VS DEPOIS)"), border=1, ln=True, align='L', fill=True)
    
    larguras = [22, 40, 40, 45, 40, 45, 45]
    colunas = ["Natureza", "Dotação Atual", "Projetado (Antes)", "Impacto Simulado", "Projetado (Depois)", "Saldo Final (Antes)", "Saldo Final (Depois)"]
    pdf.set_font("Arial", 'B', 8)
    pdf.set_fill_color(240, 240, 240)
    pdf.set_text_color(0, 0, 0)
    for i, col in enumerate(colunas):
        pdf.cell(larguras[i], 7, limpar_texto(col), border=1, align='C', fill=True)
    pdf.ln()

    pdf.set_font("Arial", '', 8)
    preenchimento = False 
    for index, row in df_comparativo.iterrows():
        if preenchimento:
            pdf.set_fill_color(245, 245, 245) 
        else:
            pdf.set_fill_color(255, 255, 255) 
            
        # AQUI FOI A CORREÇÃO: Utilizando as chaves corretas no df_comparativo
        pdf.cell(larguras[0], 6, limpar_texto(row['Natureza']), border=1, align='C', fill=preenchimento)
        pdf.cell(larguras[1], 6, limpar_texto(f_moeda_limpa(row['Dotação Atual'])), border=1, align='R', fill=preenchimento)
        pdf.cell(larguras[2], 6, limpar_texto(f_moeda_limpa(row['Projetado (Antes)'])), border=1, align='R', fill=preenchimento)
        
        if row['Impacto Simulado'] > 0:
            pdf.set_text_color(200, 0, 0)
        elif row['Impacto Simulado'] < 0:
            pdf.set_text_color(0, 150, 0)
        else:
            pdf.set_text_color(0, 0, 0)
        pdf.cell(larguras[3], 6, limpar_texto(f_moeda_limpa(row['Impacto Simulado'])), border=1, align='R', fill=preenchimento)
        
        pdf.set_text_color(0, 0, 0)
        pdf.cell(larguras[4], 6, limpar_texto(f_moeda_limpa(row['Projetado (Depois)'])), border=1, align='R', fill=preenchimento)
        
        if row['Saldo Final (Antes)'] < 0:
            pdf.set_text_color(200, 0, 0)
        else:
            pdf.set_text_color(0, 0, 0)
        pdf.cell(larguras[5], 6, limpar_texto(f_moeda_limpa(row['Saldo Final (Antes)'])), border=1, align='R', fill=preenchimento)
        
        if row['Saldo Final (Depois)'] < 0:
            pdf.set_text_color(200, 0, 0)
        else:
            pdf.set_text_color(0, 0, 0)
        pdf.cell(larguras[6], 6, limpar_texto(f_moeda_limpa(row['Saldo Final (Depois)'])), border=1, align='R', fill=preenchimento)
        
        pdf.set_text_color(0, 0, 0)
        pdf.ln()
        preenchimento = not preenchimento 

    if pdf.get_y() > 165:
        pdf.add_page() 

    pdf.ln(20)
    pdf.set_font("Arial", '', 9)
    pdf.cell(0, 5, limpar_texto("_____________________________________________________"), ln=True, align='C')
    pdf.set_font("Arial", 'B', 9)
    pdf.cell(0, 5, limpar_texto("Murilo Fontanetti"), ln=True, align='C')
    pdf.set_font("Arial", '', 9)
    pdf.cell(0, 5, limpar_texto("Contador CRC 1SP322844/0-7"), ln=True, align='C')

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf.output(tmp.name)
        with open(tmp.name, "rb") as f:
            pdf_bytes = f.read()
            
    try:
        if os.path.exists(logo_path):
            os.remove(logo_path)
    except Exception:
        pass

    return pdf_bytes

# ============================================================
# ABA 1 — PROJEÇÃO DE SALDO E FILTROS
# ============================================================
with aba1:
    st.header("1. Parâmetros e Leitura de Relatórios")
    mes_informado = st.selectbox("Mês da Última Folha Liquidada:", list(MESES.keys()))
    
    st.divider()
    st.markdown("**⚙️ Selecione as Naturezas de Despesa para análise:**")
    opcoes_naturezas = {
        "3.1.90.11": "3.1.90.11 - Vencimentos e Vantagens Fixas",
        "3.1.90.04": "3.1.90.04 - Contratação por Tempo Determinado",
        "3.1.90.13": "3.1.90.13 - Obrigações Patronais",
        "3.1.91.13": "3.1.91.13 - Obrigações Patronais (INTRA)",
        "3.1.90.16": "3.1.90.16 - Outras Despesas Variáveis",
        "3.3.90.46": "3.3.90.46 - Auxílio Alimentação"
    }

    lista_selecionadas = []
    col_chk1, col_chk2 = st.columns(2)
    for i, (codigo, descricao) in enumerate(opcoes_naturezas.items()):
        with col_chk1 if i % 2 == 0 else col_chk2:
            if st.checkbox(descricao, value=True, key=f"aba1_flag_{codigo}"):
                lista_selecionadas.append(codigo)
                
    naturezas_permitidas = tuple(lista_selecionadas)
    
    st.divider()
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        pdf_saldo = st.file_uploader("1. Saldo de Dotações", type=["pdf"])
    with col2:
        pdf_folha = st.file_uploader("2. Última Folha Liquidada", type=["pdf"])
    with col3:
        pdf_descontos_1 = st.file_uploader("3. Descontos 13º (Efetivos) [OPCIONAL]", type=["pdf"])
    with col4:
        pdf_descontos_2 = st.file_uploader("4. Descontos 13º (Temp/Pol) [OPCIONAL]", type=["pdf"])
        
    st.divider()
    if st.button("🚀 Processar Projeção", type="primary"):
        if not naturezas_permitidas:
            st.warning("⚠️ Por favor, selecione pelo menos uma Natureza de Despesa.")
        elif pdf_saldo and pdf_folha:
            with st.spinner("Lendo PDFs e calculando matemática orçamentária..."):
                
                padrao_natureza = re.compile(r'\d\.\d\.\d\d\.\d\d\.\d\d')
                padrao_valor_br = re.compile(r'\b\d{1,3}(?:\.\d{3})*,\d{2}\b|\b\d+,\d{2}\b')
                padrao_funcional = re.compile(r'\d{2}\s*\.\s*\d{3}\s*\.\s*\d{4}\s*\.\s*(\d{4})\s*\.\s*\d{4}')
                mapa_nomes = {}
                
                NATUREZAS_OBRIGATORIAS = ["3.1.90.11", "3.1.90.04", "3.1.90.13", "3.1.91.13", "3.1.90.16", "3.3.90.46"]
                
                # --- LEITURA DO PDF 1 (SALDO) ---
                dados_saldo = []
                with pdfplumber.open(pdf_saldo) as pdf:
                    codigo_atual = "0000"
                    natureza_atual = None
                    ultimo_valor = 0.0
                    
                    def arquivar_ficha_anterior():
                        if natureza_atual and (natureza_atual in NATUREZAS_OBRIGATORIAS or any(natureza_atual.startswith(n) for n in NATUREZAS_OBRIGATORIAS)):
                            dados_saldo.append({"Código": codigo_atual, "Natureza Base": natureza_atual, "Saldo com Reserva": ultimo_valor})

                    for page in pdf.pages:
                        linhas = page.extract_text().split('\n')
                        for linha in linhas:
                            if " 0000 " in linha:
                                partes = linha.split(" 0000 ")
                                num_limpo = re.sub(r'\D', '', partes[0])
                                if len(num_limpo) >= 4:
                                    codigo_atual = num_limpo[-4:]
                                ativ_nome = partes[1].strip()
                                ativ_nome = ativ_nome.replace("SIRIC", "SERIC").replace("CGEP", "SEGEP")
                                if "SESEG" in ativ_nome:
                                    ativ_nome = "Secretaria de Ordem e Segurança Pública"
                                if len(ativ_nome) > len(mapa_nomes.get(codigo_atual, "")):
                                    mapa_nomes[codigo_atual] = ativ_nome
                            
                            if re.match(r'^\s*\d{1,5}(?:\s|$)', linha):
                                arquivar_ficha_anterior()
                                natureza_atual = None
                                ultimo_valor = 0.0

                            match_nat = padrao_natureza.search(linha)
                            if match_nat:
                                arquivar_ficha_anterior()
                                natureza_atual = match_nat.group()[:9]
                                ultimo_valor = 0.0
                                
                            if natureza_atual:
                                valores = padrao_valor_br.findall(linha)
                                if valores:
                                    try:
                                        ultimo_valor = float(valores[-1].replace('.', '').replace(',', '.'))
                                    except:
                                        pass
                    arquivar_ficha_anterior()
                df_saldo = pd.DataFrame(columns=["Código", "Natureza Base", "Saldo com Reserva"]) if not dados_saldo else pd.DataFrame(dados_saldo).groupby(['Código', 'Natureza Base'], as_index=False).sum()

                # --- LEITURA DO PDF 2 (FOLHA) ---
                dados_folha = []
                with pdfplumber.open(pdf_folha) as pdf:
                    codigo_atual = "0000"
                    for page in pdf.pages:
                        for linha in page.extract_text().split('\n'):
                            if re.search(r'Proj\.*?.*Ativ', linha, re.IGNORECASE):
                                match_cod = re.search(r'\b\d{4}\b', linha)
                                if match_cod:
                                    codigo_atual = match_cod.group()
                            match_func = padrao_funcional.search(linha)
                            if match_func:
                                codigo_atual = match_func.group(1) 
                            match = padrao_natureza.search(linha)
                            if match:
                                nat_det = match.group()
                                nat_base = nat_det[:9]
                                if nat_base in NATUREZAS_OBRIGATORIAS or any(nat_det.startswith(n) for n in NATUREZAS_OBRIGATORIAS):
                                    is_13 = nat_det in ["3.1.90.11.43", "3.1.90.04.13"]
                                    is_base_patronal = nat_det in ["3.1.90.11.74", "3.1.90.11.75"] or nat_base == "3.1.90.04"
                                    is_base_inss_ferias = nat_det in ["3.1.90.11.74", "3.1.90.11.75"] or nat_base == "3.1.90.04"
                                    try:
                                        v_num = float(linha.split()[-4].replace('.', '').replace(',', '.'))
                                        dados_folha.append({"Código": codigo_atual, "Natureza Base": nat_base, "Liquidado": v_num, "Embutido 13": v_num if is_13 else 0.0, "Base Patronal": v_num if is_base_patronal else 0.0, "Base INSS Ferias": v_num if is_base_inss_ferias else 0.0})
                                    except:
                                        pass
                df_folha = pd.DataFrame(columns=["Código", "Natureza Base", "Liquidado", "Embutido 13", "Base Patronal", "Base INSS Ferias"]) if not dados_folha else pd.DataFrame(dados_folha).groupby(['Código', 'Natureza Base'], as_index=False).sum()

                # --- LEITURA DOS PDFs 3 E 4 (DESCONTOS MÚLTIPLOS) ---
                dados_desc = []
                def extrair_descontos(pdf_file):
                    if pdf_file is not None:
                        with pdfplumber.open(pdf_file) as pdf:
                            codigo_atual_desc = "0000"
                            for page in pdf.pages:
                                for linha in page.extract_text().split('\n'):
                                    if re.search(r'Proj\.*?.*Ativ', linha, re.IGNORECASE):
                                        match_cod = re.search(r'\b\d{4}\b', linha)
                                        if match_cod:
                                            codigo_atual_desc = match_cod.group()
                                    match_func = padrao_funcional.search(linha)
                                    if match_func:
                                        codigo_atual_desc = match_func.group(1)
                                    match_nat = padrao_natureza.search(linha)
                                    if match_nat:
                                        nat_det = match_nat.group()
                                        nat_base = nat_det[:9]
                                        if nat_base in NATUREZAS_OBRIGATORIAS or any(nat_det.startswith(n) for n in NATUREZAS_OBRIGATORIAS):
                                            if nat_base == '3.1.90.04' and nat_det != '3.1.90.04.13':
                                                continue 
                                            try:
                                                v_num = float(linha.split()[-4].replace('.', '').replace(',', '.'))
                                                if v_num > 0:
                                                    dados_desc.append({"Código": codigo_atual_desc, "Natureza Base": nat_base, "Desconto 13": v_num})
                                            except:
                                                pass
                
                extrair_descontos(pdf_descontos_1)
                extrair_descontos(pdf_descontos_2)
                
                df_desc = pd.DataFrame(columns=["Código", "Natureza Base", "Desconto 13"]) if not dados_desc else pd.DataFrame(dados_desc).groupby(['Código', 'Natureza Base'], as_index=False).sum()

                # --- SÍNTESE E MATEMÁTICA ---
                if df_saldo.empty and df_folha.empty:
                    st.warning("⚠️ Nenhuma natureza válida foi encontrada nos PDFs selecionados.")
                else:
                    df_merge = pd.merge(df_saldo, df_folha, on=['Código', 'Natureza Base'], how='outer')
                    df_merge = pd.merge(df_merge, df_desc, on=['Código', 'Natureza Base'], how='outer')
                    df_merge.fillna(0, inplace=True)
                    df_merge = df_merge[(df_merge['Saldo com Reserva'] > 0) | (df_merge['Liquidado'] > 0)].copy()

                    mes_rest = 12 - MESES[mes_informado]
                    mes_calc = max(mes_rest, 1)

                    df_merge['Salário Limpo Mensal'] = df_merge['Liquidado'] - df_merge['Embutido 13']
                    df_merge['Projeção Salarial'] = df_merge['Salário Limpo Mensal'] * mes_rest
                    df_merge['Base 13 Bruta'] = (df_merge['Salário Limpo Mensal'] - df_merge['Desconto 13']).clip(lower=0)

                    df_merge['Base Limpa Patronal'] = df_merge.apply(lambda r: r['Base Patronal'] - r['Embutido 13'] if r['Natureza Base'] == '3.1.90.04' else r['Base Patronal'], axis=1)
                    df_merge['Base Limpa INSS Ferias'] = df_merge.apply(lambda r: r['Base INSS Ferias'] - r['Embutido 13'] if r['Natureza Base'] == '3.1.90.04' else r['Base INSS Ferias'], axis=1)
                    
                    base_patronal_codigo = df_merge.groupby('Código')['Base Limpa Patronal'].sum().to_dict()
                    base_ferias_codigo = df_merge.groupby('Código')['Base Limpa INSS Ferias'].sum().to_dict()

                    def aplicar_13(row):
                        nb = row['Natureza Base']
                        if nb in ['3.3.90.46', '3.1.90.16']:
                            return 0.0
                        elif nb == '3.1.90.13':
                            return max(0.0, (base_patronal_codigo.get(row['Código'], 0.0) * 0.22) - row['Desconto 13'])
                        else:
                            return row['Base 13 Bruta']

                    def aplicar_ferias(row):
                        nb = row['Natureza Base']
                        if nb in ['3.1.90.11', '3.1.90.04']:
                            return ((row['Salário Limpo Mensal'] / 3) / 12) * mes_rest
                        elif nb == '3.1.90.13':
                            return max(0.0, ((base_ferias_codigo.get(row['Código'], 0.0) / 3) / 12) * mes_rest * 0.22)
                        return 0.0

                    df_merge['Provisão 13º'] = df_merge.apply(aplicar_13, axis=1)
                    df_merge['Provisão Férias'] = df_merge.apply(aplicar_ferias, axis=1)
                    df_merge['Provisão 13º'] = df_merge.apply(lambda r: 0.0 if r['Natureza Base'] in ['3.3.90.46', '3.1.90.16'] else r['Provisão 13º'], axis=1)
                    df_merge['Provisão Férias'] = df_merge.apply(lambda r: 0.0 if r['Natureza Base'] in ['3.3.90.46', '3.1.90.16'] else r['Provisão Férias'], axis=1)
                    
                    df_merge['Total Projetado'] = df_merge['Projeção Salarial'] + df_merge['Provisão 13º'] + df_merge['Provisão Férias']
                    df_merge['Saldo Final'] = df_merge['Saldo com Reserva'] - df_merge['Total Projetado']
                    
                    def identificar_mes_credito(row):
                        if row['Saldo Final'] >= 0 or row['Total Projetado'] <= 0:
                            return "✅ Suficiente"
                        custo_mensal = row['Total Projetado'] / mes_calc
                        if custo_mensal == 0:
                            return "✅ Suficiente"
                        mes_falta = MESES[mes_informado] + int(row['Saldo com Reserva'] / custo_mensal) + 1
                        return f"⚠️ Adicionar Crédito em {num_para_mes[mes_falta]}" if mes_falta <= 12 else "✅ Suficiente"

                    df_merge['Mês do Crédito'] = df_merge.apply(identificar_mes_credito, axis=1)
                    df_merge['Atividade'] = df_merge['Código'].map(mapa_nomes).fillna("—")
                    
                    df_merge = df_merge[df_merge['Natureza Base'].isin(naturezas_permitidas)].copy()
                    
                    df_merge['Filtro Visual'] = df_merge['Código'] + " - " + df_merge['Atividade']
                    df_merge['% Saldo'] = (df_merge['Saldo Final'] / df_merge['Saldo com Reserva'].replace(0, 1)) * 100
                    df_merge.sort_values('% Saldo', ascending=True, inplace=True)
                    
                    st.session_state.df_processado = df_merge
                    st.success("✅ Relatórios processados! Desça a página para ver o Dashboard e as Exportações.")

        else:
            st.error("⚠️ Por favor, anexe pelo menos os relatórios de Saldo e da Folha para calcular.")

    # =========================================================
    # PAINEL DE EXIBIÇÃO DINÂMICO E EXPORTAÇÕES (ABA 1)
    # =========================================================
    if st.session_state.df_processado is not None:
        df_final = st.session_state.df_processado.copy()
        
        st.divider()
        st.header("2. Dashboard de Resultados e Exportações")
        
        st.info("👇 **USE A CAIXA ABAIXO PARA ISOLAR UMA ATIVIDADE ESPECÍFICA:**")
        lista_opcoes = ["Todas as atividades"] + sorted(df_final['Filtro Visual'].unique().tolist())
        filtro_selecionado = st.selectbox("Selecione a Atividade para análise:", lista_opcoes)
        
        if filtro_selecionado != "Todas as atividades":
            df_final = df_final[df_final['Filtro Visual'] == filtro_selecionado]

        COLS_RELATORIO = ['Código', 'Atividade', 'Natureza Base', 'Saldo com Reserva', 'Liquidado', 'Projeção Salarial', 'Provisão 13º', 'Provisão Férias', 'Total Projetado', 'Saldo Final', 'Mês do Crédito']
        df_exp = df_final[COLS_RELATORIO].copy()
        df_exp.columns = ['Cód', 'Atividade', 'Natureza', 'Saldo Atual', 'Liquidado', 'Proj. Sal.', 'Prov. 13º', 'Prov. Férias', 'Total Projet.', 'Saldo Final', 'Situação']

        st.markdown("### 📥 Escolha o Formato do Documento")
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        with col_btn1:
            st.download_button("📄 Baixar PDF Oficial", gerar_pdf(df_final, mes_informado), f"Projecao_{mes_informado}.pdf", "application/pdf", use_container_width=True)
        with col_btn2:
            st.download_button("📝 Baixar em Word", gerar_word_aba1(df_exp, mes_informado), f"Projecao_{mes_informado}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with col_btn3:
            st.download_button("📊 Baixar em Excel", gerar_excel_aba1(df_exp), f"Projecao_{mes_informado}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            
        st.divider()

        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Total Saldo Atual", fmt_br(df_final['Saldo com Reserva'].sum()))
        col_m2.metric("Total Projetado", fmt_br(df_final['Total Projetado'].sum()))
        col_m3.metric("Saldo Projetado Final", fmt_br(df_final['Saldo Final'].sum()))
        
        st.divider()
        insuf = df_final[df_final['Saldo Final'] < 0]
        sufic = df_final[df_final['Saldo Final'] >= 0]
        
        if not insuf.empty:
            st.subheader("🔴 Relatório de Insuficiência de Saldo")
            st.dataframe(insuf[COLS_RELATORIO].round(2), use_container_width=True)
        else:
            st.info("✅ Nenhuma ficha com saldo insuficiente encontrada nesta seleção.")
            
        st.subheader("🟢 Relatório de Suficiência de Saldo")
        st.dataframe(sufic[COLS_RELATORIO].round(2), use_container_width=True)
        
        st.divider()
        st.subheader("📊 Resumo por Atividade")
        resumo = df_final.groupby(['Código', 'Atividade']).agg(Saldo_Total=('Saldo com Reserva', 'sum'), Total_Projetado=('Total Projetado', 'sum'), Saldo_Final=('Saldo Final', 'sum')).reset_index()
        resumo['Situação'] = resumo['Saldo_Final'].apply(lambda x: "✅ Suficiente" if x >= 0 else "🔴 Insuficiente")
        st.dataframe(resumo.round(2), use_container_width=True)

# ============================================================
# ABA 2 — SIMULADOR DE ACRÉSCIMO / DECRÉSCIMO COM FLAGS E MISTO
# ============================================================
with aba2:
    st.header("👤 Simulador Integrado com Seleção de Naturezas")
    
    if st.session_state.df_processado is None:
        st.warning("⚠️ O Simulador precisa da base orçamentária atualizada. Por favor, vá até a aba 'Análise de Folha' e processe os relatórios primeiro.")
    else:
        st.markdown("Marque com **Flags** (✓) as naturezas de despesa que farão parte desta análise de impacto e informe os valores.")

        c1, c2 = st.columns([1, 1])
        with c1:
            lista_atividades_simulador = sorted(st.session_state.df_processado['Filtro Visual'].unique().tolist())
            sim_atividade_visual = st.selectbox("1. Atividade que sofrerá a alteração:", lista_atividades_simulador)
            sim_operacao = st.radio("2. Operação:", ["➕ Acréscimo", "➖ Decréscimo", "🔀 Mista (Acréscimo e Decréscimo)"])
            sim_mes = st.selectbox("3. Mês de início da mudança:", list(MESES.keys()), key="sim_mes")
            sim_descricao = st.text_input("4. Descrição da Estimativa (Opcional):", placeholder="Ex: Contratação de 2 Professores e Exoneração de 1", key="sim_desc")

        with c2:
            st.markdown("**5. Selecione as Naturezas (Flags) e insira o valor mensal:**")
            
            flag_11 = st.checkbox("3.1.90.11 - Vencimentos e Vantagens Fixas", value=True, key="aba2_chk_11")
            sim_11 = 0.0
            if flag_11:
                if "Mista" in sim_operacao:
                    colA, colB = st.columns(2)
                    with colA: v_in11 = st.number_input("➕ Entra (R$):", min_value=0.0, step=100.0, key="a_11_in")
                    with colB: v_out11 = st.number_input("➖ Sai (R$):", min_value=0.0, step=100.0, key="a_11_out")
                    sim_11 = v_in11 - v_out11
                else:
                    val_11 = st.number_input("Valor Mensal (3.1.90.11):", min_value=0.0, step=100.0, key="a_num_11")
                    sim_11 = val_11 if "Acréscimo" in sim_operacao else -val_11

            flag_04 = st.checkbox("3.1.90.04 - Contratação por Tempo Determinado", key="aba2_chk_04")
            sim_04 = 0.0
            if flag_04:
                if "Mista" in sim_operacao:
                    colA, colB = st.columns(2)
                    with colA: v_in04 = st.number_input("➕ Entra (R$):", min_value=0.0, step=100.0, key="a_04_in")
                    with colB: v_out04 = st.number_input("➖ Sai (R$):", min_value=0.0, step=100.0, key="a_04_out")
                    sim_04 = v_in04 - v_out04
                else:
                    val_04 = st.number_input("Valor Mensal (3.1.90.04):", min_value=0.0, step=100.0, key="a_num_04")
                    sim_04 = val_04 if "Acréscimo" in sim_operacao else -val_04

            flag_16 = st.checkbox("3.1.90.16 - Outras Despesas Variáveis", key="aba2_chk_16")
            sim_16 = 0.0
            if flag_16:
                if "Mista" in sim_operacao:
                    colA, colB = st.columns(2)
                    with colA: v_in16 = st.number_input("➕ Entra (R$):", min_value=0.0, step=50.0, key="a_16_in")
                    with colB: v_out16 = st.number_input("➖ Sai (R$):", min_value=0.0, step=50.0, key="a_16_out")
                    sim_16 = v_in16 - v_out16
                else:
                    val_16 = st.number_input("Valor Mensal (3.1.90.16):", min_value=0.0, step=50.0, key="a_num_16")
                    sim_16 = val_16 if "Acréscimo" in sim_operacao else -val_16

            flag_46 = st.checkbox("3.3.90.46 - Auxílio Alimentação", key="aba2_chk_46")
            sim_46 = 0.0
            if flag_46:
                if "Mista" in sim_operacao:
                    colA, colB = st.columns(2)
                    with colA: v_in46 = st.number_input("➕ Entra (R$):", min_value=0.0, step=50.0, key="a_46_in")
                    with colB: v_out46 = st.number_input("➖ Sai (R$):", min_value=0.0, step=50.0, key="a_46_out")
                    sim_46 = v_in46 - v_out46
                else:
                    val_46 = st.number_input("Valor Mensal (3.3.90.46):", min_value=0.0, step=50.0, key="a_num_46")
                    sim_46 = val_46 if "Acréscimo" in sim_operacao else -val_46
            
            st.markdown("**6. Encargos Patronais (Cálculo Automático 22%):**")
            flag_9113 = st.checkbox("3.1.91.13 - RPPS Próprio (Isento nas Férias)", key="aba2_chk_9113")
            flag_9013 = st.checkbox("3.1.90.13 - INSS Geral (Incide nas Férias)", key="aba2_chk_9013")

        if st.button("📊 Processar Simulação Modular", type="primary"):
            codigo_selecionado = sim_atividade_visual.split(" - ")[0]
            meses_impacto = 13 - MESES[sim_mes] 
            
            dict_impactos = {}
            dados_simulacao = []
            
            if flag_11 and sim_11 != 0:
                imp_11 = (sim_11 * meses_impacto + (sim_11 / 12) * meses_impacto + ((sim_11 / 3) / 12) * meses_impacto)
                dict_impactos["3.1.90.11"] = imp_11
                dados_simulacao.append({"Item": "3.1.90.11 (Salário + 13º + Férias)", "Valor": imp_11})
                
            if flag_04 and sim_04 != 0:
                imp_04 = (sim_04 * meses_impacto + (sim_04 / 12) * meses_impacto + ((sim_04 / 3) / 12) * meses_impacto)
                dict_impactos["3.1.90.04"] = imp_04
                dados_simulacao.append({"Item": "3.1.90.04 (Salário + 13º + Férias)", "Valor": imp_04})
                
            if flag_16 and sim_16 != 0:
                imp_16 = sim_16 * meses_impacto
                dict_impactos["3.1.90.16"] = imp_16
                dados_simulacao.append({"Item": "3.1.90.16 (Despesa Variável Mensal)", "Valor": imp_16})
                
            if flag_46 and sim_46 != 0:
                imp_46 = sim_46 * meses_impacto
                dict_impactos["3.3.90.46"] = imp_46
                dados_simulacao.append({"Item": "3.3.90.46 (Auxílio Alimentação)", "Valor": imp_46})
                
            base_salario_mensal = (sim_11 + sim_04) * meses_impacto
            base_13 = ((sim_11 + sim_04) / 12) * meses_impacto
            base_ferias = (((sim_11 + sim_04) / 3) / 12) * meses_impacto
            
            if flag_9113:
                imp_9113 = (base_salario_mensal + base_13) * 0.22
                dict_impactos["3.1.91.13"] = imp_9113
                if imp_9113 != 0:
                    dados_simulacao.append({"Item": "3.1.91.13 (Encargo RPPS 22% - S/ Férias)", "Valor": imp_9113})
                    
            if flag_9013:
                imp_9013 = (base_salario_mensal + base_13 + base_ferias) * 0.22
                dict_impactos["3.1.90.13"] = imp_9013
                if imp_9013 != 0:
                    dados_simulacao.append({"Item": "3.1.90.13 (Encargo INSS 22% - C/ Férias)", "Valor": imp_9013})

            impacto_final = sum(dict_impactos.values())
            df_simulacao = pd.DataFrame(dados_simulacao)

            if not dict_impactos:
                st.warning("⚠️ Selecione pelo menos uma natureza de despesa e informe um valor diferente de zero para simular.")
            else:
                df_atual = st.session_state.df_processado
                df_atividade = df_atual[df_atual['Código'] == codigo_selecionado].copy()
                
                todas_naturezas = set(df_atividade['Natureza Base'].tolist() + list(dict_impactos.keys()))
                dados_comp = []
                
                for nat in todas_naturezas:
                    row_nat = df_atividade[df_atividade['Natureza Base'] == nat]
                    saldo_dot = row_nat['Saldo com Reserva'].sum() if not row_nat.empty else 0.0
                    proj_antes = row_nat['Total Projetado'].sum() if not row_nat.empty else 0.0
                    final_antes = row_nat['Saldo Final'].sum() if not row_nat.empty else 0.0
                    
                    imp = dict_impactos.get(nat, 0.0)
                    proj_depois = proj_antes + imp
                    final_depois = saldo_dot - proj_depois
                    
                    if saldo_dot > 0 or imp != 0 or proj_antes > 0:
                        dados_comp.append({
                            "Natureza": nat, "Dotação Atual": saldo_dot, "Projetado (Antes)": proj_antes,
                            "Impacto Simulado": imp, "Projetado (Depois)": proj_depois, "Saldo Final (Antes)": final_antes, "Saldo Final (Depois)": final_depois
                        })
                
                df_comparativo = pd.DataFrame(dados_comp).sort_values("Natureza")
                
                st.divider()
                st.subheader(f"⚖️ Quadro Comparativo — {sim_atividade_visual}")
                
                col_res1, col_res2 = st.columns([2, 1])
                with col_res1:
                    st.markdown("**Memória de Cálculo da Simulação:**")
                    st.dataframe(df_simulacao.style.format({"Valor": "R$ {:,.2f}"}), hide_index=True, use_container_width=True)
                with col_res2:
                    status_cor = "🟢 Economia de" if impacto_final < 0 else "🔴 Custo de"
                    st.metric(label=f"Impacto Financeiro Global ({status_cor})", value=fmt_br(abs(impacto_final)))

                st.markdown("**Quadro Geral da Dotação (Antes vs Depois):**")
                st.dataframe(
                    df_comparativo.style.format({
                        "Dotação Atual": "R$ {:,.2f}", "Projetado (Antes)": "R$ {:,.2f}", "Impacto Simulado": "R$ {:,.2f}", 
                        "Projetado (Depois)": "R$ {:,.2f}", "Saldo Final (Antes)": "R$ {:,.2f}", "Saldo Final (Depois)": "R$ {:,.2f}"
                    }).map(
                        lambda x: 'color: red;' if x > 0 else ('color: green;' if x < 0 else ''), subset=['Impacto Simulado']
                    ).map(
                        lambda x: 'color: red;' if x < 0 else 'color: green;', subset=['Saldo Final (Antes)', 'Saldo Final (Depois)']
                    ), 
                    hide_index=True, use_container_width=True
                )
                
                st.divider()
                st.markdown("### 📥 Escolha o Formato da Simulação")
                col_btnA, col_btnB, col_btnC = st.columns(3)
                
                with col_btnA:
                    st.download_button("📄 Baixar em PDF", gerar_pdf_simulador(sim_atividade_visual, sim_operacao, sim_mes, impacto_final, df_comparativo, df_simulacao, sim_descricao), f"Simulacao_{codigo_selecionado}.pdf", "application/pdf", use_container_width=True)
                with col_btnB:
                    st.download_button("📝 Baixar em Word", gerar_word_simulador(sim_atividade_visual, sim_operacao, sim_mes, impacto_final, df_comparativo, df_simulacao, sim_descricao), f"Simulacao_{codigo_selecionado}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                with col_btnC:
                    st.download_button("📊 Baixar em Excel", gerar_excel_simulador(df_comparativo, df_simulacao), f"Simulacao_{codigo_selecionado}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
